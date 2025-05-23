import os
import re
import platform
import docx
import PyPDF2 

# Attempt to import pywin32 for .doc processing on Windows
IS_WINDOWS = platform.system() == "Windows"
word_app = None
if IS_WINDOWS:
    try:
        import win32com.client as win32
        # Try to create a Word application instance once
        try:
             word_app = win32.Dispatch("Word.Application")
             word_app.Visible = False # Keep Word hidden
             print("Microsoft Word instance created successfully for .doc processing.")
        except Exception as word_init_e:
             print(f"Warning: Could not initialize Microsoft Word COM object: {word_init_e}")
             print("  .doc file processing will be skipped. Ensure Word is installed.")
             word_app = None # Ensure it's None if initialization failed
    except ImportError:
        print("Warning: 'pywin32' library not found.")
        print("  Install it using 'pip install pywin32' to enable .doc processing on Windows.")
        print("  .doc files will be skipped.")
    except Exception as e:
        print(f"An unexpected error occurred during pywin32 import or Word initialization: {e}")
        print("  .doc file processing will be skipped.")


# --- Configuration ---
INPUT_FOLDER = 'Resumes_Docx'
OUTPUT_FOLDER = 'Processed_Resumes_Txt'
# --- End Configuration ---

def clean_filename_part(name):
    """Removes characters invalid for file/folder names and cleans up."""
    name = re.sub(r'[<>:"/\\|?*]', '_', name)
    name = name.strip().strip('.')
    return name

def read_docx(file_path):
    """Reads text content from a .docx file."""
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                         full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"  Error reading DOCX {os.path.basename(file_path)}: {e}")
        return None

def read_pdf(file_path):
    """Reads text content from a .pdf file."""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            num_pages = len(reader.pages)
            for page_num in range(num_pages):
                try:
                    page = reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                         text += page_text + "\n"
                except Exception as page_e:
                    print(f"  Warning: Could not extract text from page {page_num + 1} of PDF {os.path.basename(file_path)}: {page_e}")
                    continue
        return text.strip() if text else None
    except Exception as e:
        print(f"  Error reading PDF {os.path.basename(file_path)}: {e}")
        return None

def read_doc_windows(file_path):
    """Reads text content from a .doc file using pywin32 (Windows only)."""
    global word_app # Use the global word_app instance
    if not word_app:
        print(f"  Skipping .doc file (Word COM object not available): {os.path.basename(file_path)}")
        return None

    doc = None # Initialize doc to None
    try:
        # Use os.path.abspath to ensure the path format is correct for COM
        abs_path = os.path.abspath(file_path)
        # Open the document
        doc = word_app.Documents.Open(abs_path, ReadOnly=True)
        # Extract text
        content = doc.Content.Text
        return content.strip() if content else None
    except Exception as e:
        print(f"  Error reading DOC {os.path.basename(file_path)} with pywin32: {e}")
        return None
    finally:
        # Close the document without saving changes
        if doc:
            try:
                doc.Close(SaveChanges=0) # 0 means do not save changes
            except Exception as close_e:
                print(f"  Error closing Word document {os.path.basename(file_path)}: {close_e}")


# --- Main Script ---
if not os.path.isdir(INPUT_FOLDER):
    print(f"Error: Input folder '{INPUT_FOLDER}' not found.")
    print("Please make sure the script is in the same directory as the 'Resumes_Docx' folder.")
else:
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    print(f"Created/Ensured output folder: '{OUTPUT_FOLDER}'")

    processed_files = 0
    failed_files = 0
    skipped_files = 0

    for subdir, dirs, files in os.walk(INPUT_FOLDER):
        for filename in files:
            original_file_path = os.path.join(subdir, filename)
            base_name, extension = os.path.splitext(filename)
            extension = extension.lower()

            relative_subdir = os.path.relpath(subdir, INPUT_FOLDER)
            original_subfolder_name = os.path.basename(relative_subdir) if relative_subdir != '.' else 'Root'
            cleaned_subfolder_name = clean_filename_part(original_subfolder_name)
            cleaned_base_name = clean_filename_part(base_name)

            print(f"\nProcessing: {original_file_path}")

            content = None
            if extension == '.docx':
                content = read_docx(original_file_path)
            elif extension == '.pdf':
                content = read_pdf(original_file_path)
            elif extension == '.doc':
                if IS_WINDOWS and word_app: # Check if on Windows and Word object is ready
                    content = read_doc_windows(original_file_path)
                else:
                    if not IS_WINDOWS:
                        print(f"  Skipping .doc file (OS is not Windows): {filename}.")
                    elif not word_app:
                         print(f"  Skipping .doc file (Word/pywin32 issue): {filename}.")
                    else: # Should not happen based on checks, but just in case
                         print(f"  Skipping .doc file: {filename}.")
                    skipped_files += 1
                    continue
            else:
                if filename != os.path.basename(__file__) and extension not in ['.py', '.csv', '.txt', '.md']:
                     print(f"  Skipping file with unsupported extension: {filename}")
                skipped_files += 1
                continue

            if content is not None and content.strip():
                output_txt_filename = f"{cleaned_subfolder_name}-{cleaned_base_name}.txt"
                output_txt_path = os.path.join(OUTPUT_FOLDER, output_txt_filename)
                try:
                    with open(output_txt_path, 'w', encoding='utf-8') as txt_file:
                        txt_file.write(content)
                    print(f"  Successfully saved text to: {output_txt_path}")
                    processed_files += 1
                except Exception as e:
                    print(f"  Error writing text file {output_txt_path}: {e}")
                    failed_files += 1
            elif content is None:
                 print(f"  Failed to read content from {filename}.")
                 failed_files += 1
            else:
                 print(f"  Skipping {filename} due to empty extracted content.")
                 failed_files += 1


    print(f"\n--- Processing Complete ---")
    print(f"Successfully processed: {processed_files} files")
    print(f"Failed reading/writing content: {failed_files} files")
    print(f"Skipped (unsupported/doc/other): {skipped_files} files")

    # Clean up the Word application instance if it was created
    if word_app:
        try:
            word_app.Quit()
            print("\nClosed Microsoft Word instance.")
        except Exception as quit_e:
            print(f"Warning: Error quitting Word application: {quit_e}")

